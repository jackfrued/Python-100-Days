"""
读取Word文件

Version: 0.1
Author: 骆昊
Date: 2018-03-26
"""

from docx import Document

doc = Document('./res/用函数还是用复杂的表达式.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].runs[0].text)

content = []
for para in doc.paragraphs:
    content.append(para.text)
print(''.join(content))
